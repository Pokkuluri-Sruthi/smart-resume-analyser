import re
from PyPDF2 import PdfReader
from docx import Document

SKILLS = [
    "python","java","c","c++","javascript","html","css","react","node.js",
    "node","express","sql","mysql","postgresql","mongodb","flask","django",
    "machine learning","deep learning","artificial intelligence","data analysis",
    "pandas","numpy","matplotlib","power bi","excel","tableau","git","github",
    "docker","aws","azure","linux","tensorflow","pytorch","spark","pyspark",
    "communication","leadership"
]

def extract_pdf_text(filepath):
    text = ""
    reader = PdfReader(filepath)
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_docx_text(filepath):import re
from PyPDF2 import PdfReader
from docx import Document


SKILLS = [
    "python",
    "java",
    "c",
    "c++",
    "javascript",
    "html",
    "css",
    "react",
    "node.js",
    "node",
    "express",
    "sql",
    "mysql",
    "postgresql",
    "mongodb",
    "flask",
    "django",
    "machine learning",
    "deep learning",
    "artificial intelligence",
    "data analysis",
    "pandas",
    "numpy",
    "matplotlib",
    "power bi",
    "excel",
    "tableau",
    "git",
    "github",
    "docker",
    "aws",
    "azure",
    "linux",
    "tensorflow",
    "pytorch",
    "spark",
    "pyspark",
    "communication",
    "leadership"
]


def extract_pdf_text(filepath):

    text = ""

    reader = PdfReader(filepath)

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def extract_docx_text(filepath):

    document = Document(filepath)

    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    return text


def extract_text(filepath):

    if filepath.lower().endswith(".pdf"):
        return extract_pdf_text(filepath)

    if filepath.lower().endswith(".docx"):
        return extract_docx_text(filepath)

    return ""


def find_skills(text):

    text = text.lower()

    found_skills = []

    for skill in SKILLS:

        if skill.lower() in text:
            found_skills.append(skill)

    return sorted(set(found_skills))


def extract_email(text):

    pattern = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"

    match = re.search(pattern, text)

    if match:
        return match.group()

    return "Not detected"


def extract_phone(text):

    pattern = r"\+?\d[\d\s\-]{8,14}\d"

    match = re.search(pattern, text)

    if match:
        return match.group()

    return "Not detected"


def detect_sections(text):

    text_lower = text.lower()

    sections = {
        "Education": [
            "education",
            "academic"
        ],

        "Experience": [
            "experience",
            "work experience",
            "employment"
        ],

        "Projects": [
            "projects",
            "project"
        ],

        "Certifications": [
            "certification",
            "certifications"
        ],

        "Skills": [
            "skills",
            "technical skills"
        ]
    }

    result = {}

    for section, keywords in sections.items():

        result[section] = any(
            keyword in text_lower
            for keyword in keywords
        )

    return result


def calculate_score(
    text,
    skills,
    sections,
    job_description
):

    score = 0

    # Resume length
    words = len(text.split())

    if words >= 300:
        score += 20

    elif words >= 150:
        score += 15

    else:
        score += 8

    # Skills
    if len(skills) >= 10:
        score += 25

    elif len(skills) >= 5:
        score += 18

    elif len(skills) >= 2:
        score += 10

    # Sections
    section_count = sum(sections.values())

    score += min(section_count * 5, 25)

    # Contact information
    if extract_email(text) != "Not detected":
        score += 5

    if extract_phone(text) != "Not detected":
        score += 5

    # Job description
    if job_description.strip():

        job_skills = find_skills(job_description)

        if job_skills:

            matched = set(skills) & set(job_skills)

            match_percentage = (
                len(matched) / len(job_skills)
            ) * 20

            score += int(match_percentage)

    return min(score, 100)


def analyze_resume(filepath, job_description=""):

    text = extract_text(filepath)

    if not text.strip():

        return {
            "score": 0,
            "email": "Not detected",
            "phone": "Not detected",
            "skills": [],
            "matched_skills": [],
            "missing_skills": [],
            "sections": {},
            "suggestions": [
                "Could not extract text from the resume."
            ]
        }

    skills = find_skills(text)

    sections = detect_sections(text)

    email = extract_email(text)

    phone = extract_phone(text)

    job_skills = find_skills(job_description)

    matched_skills = sorted(
        set(skills) & set(job_skills)
    )

    missing_skills = sorted(
        set(job_skills) - set(skills)
    )

    score = calculate_score(
        text,
        skills,
        sections,
        job_description
    )

    suggestions = []

    if len(skills) < 5:
        suggestions.append(
            "Add more relevant technical skills."
        )

    if not sections["Projects"]:
        suggestions.append(
            "Add a Projects section with 2-3 strong projects."
        )

    if not sections["Experience"]:
        suggestions.append(
            "Add internship, training or experience details."
        )

    if not sections["Certifications"]:
        suggestions.append(
            "Add relevant certifications if available."
        )

    if email == "Not detected":
        suggestions.append(
            "Add a professional email address."
        )

    if phone == "Not detected":
        suggestions.append(
            "Add a contact phone number."
        )

    if job_description and missing_skills:

        suggestions.append(
            "Consider learning or highlighting: "
            + ", ".join(missing_skills)
        )

    if not suggestions:

        suggestions.append(
            "Your resume has a good overall structure."
        )

    return {
        "score": score,
        "email": email,
        "phone": phone,
        "skills": skills,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "sections": sections,
        "suggestions": suggestions
    }
    document = Document(filepath)
    return "\n".join(p.text for p in document.paragraphs)

def extract_text(filepath):
    if filepath.lower().endswith(".pdf"):
        return extract_pdf_text(filepath)
    if filepath.lower().endswith(".docx"):
        return extract_docx_text(filepath)
    return ""

def find_skills(text):
    text = text.lower()
    return sorted(set(skill for skill in SKILLS if skill.lower() in text))

def extract_email(text):
    match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    return match.group() if match else "Not detected"

def extract_phone(text):
    match = re.search(r"\+?\d[\d\s\-]{8,14}\d", text)
    return match.group() if match else "Not detected"

def detect_sections(text):
    text_lower = text.lower()
    sections = {
        "Education": ["education", "academic"],
        "Experience": ["experience", "work experience", "employment"],
        "Projects": ["projects", "project"],
        "Certifications": ["certification", "certifications"],
        "Skills": ["skills", "technical skills"]
    }
    return {section: any(k in text_lower for k in keywords) for section, keywords in sections.items()}

def calculate_score(text, skills, sections, job_description):
    score = 0
    words = len(text.split())
    score += 20 if words >= 300 else 15 if words >= 150 else 8
    score += 25 if len(skills) >= 10 else 18 if len(skills) >= 5 else 10 if len(skills) >= 2 else 0
    score += min(sum(sections.values()) * 5, 25)
    if extract_email(text) != "Not detected":
        score += 5
    if extract_phone(text) != "Not detected":
        score += 5
    if job_description.strip():
        job_skills = find_skills(job_description)
        if job_skills:
            matched = set(skills) & set(job_skills)
            score += int((len(matched) / len(job_skills)) * 20)
    return min(score, 100)

def analyze_resume(filepath, job_description=""):
    text = extract_text(filepath)
    if not text.strip():
        return {
            "score": 0, "email": "Not detected", "phone": "Not detected",
            "skills": [], "matched_skills": [], "missing_skills": [],
            "sections": {}, "suggestions": ["Could not extract text from the resume."]
        }

    skills = find_skills(text)
    sections = detect_sections(text)
    email = extract_email(text)
    phone = extract_phone(text)
    job_skills = find_skills(job_description)
    matched_skills = sorted(set(skills) & set(job_skills))
    missing_skills = sorted(set(job_skills) - set(skills))
    score = calculate_score(text, skills, sections, job_description)

    suggestions = []
    if len(skills) < 5:
        suggestions.append("Add more relevant technical skills.")
    if not sections["Projects"]:
        suggestions.append("Add a Projects section with 2-3 strong projects.")
    if not sections["Experience"]:
        suggestions.append("Add internship, training or experience details.")
    if not sections["Certifications"]:
        suggestions.append("Add relevant certifications if available.")
    if email == "Not detected":
        suggestions.append("Add a professional email address.")
    if phone == "Not detected":
        suggestions.append("Add a contact phone number.")
    if job_description and missing_skills:
        suggestions.append("Consider learning or highlighting: " + ", ".join(missing_skills))
    if not suggestions:
        suggestions.append("Your resume has a good overall structure.")

    return {
        "score": score, "email": email, "phone": phone, "skills": skills,
        "matched_skills": matched_skills, "missing_skills": missing_skills,
        "sections": sections, "suggestions": suggestions
    }
